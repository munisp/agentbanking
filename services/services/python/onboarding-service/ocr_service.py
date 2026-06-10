import sys as _sys, os as _os
_sys.path.insert(0, _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), ".."))
from shared.middleware import apply_middleware, ErrorResponse
from shared.observability import setup_logging, get_logger, metrics_router, MetricsMiddleware
"""
Production-Ready OCR Service
Multi-engine pipeline using PaddleOCR, VLM, and Docling for document processing
Integrates with: PostgreSQL, Kafka, Redis, Lakehouse
"""

import os
import uuid
import logging
import json
import hashlib
import base64
import tempfile
import re
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any, AsyncGenerator, Tuple
from contextlib import asynccontextmanager
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
import io

import asyncpg
import redis.asyncio as redis
from fastapi import FastAPI, HTTPException, Depends, File, UploadFile, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware

apply_middleware(app)
setup_logging("ocr-service-(production)")
app.include_router(metrics_router)

from pydantic import BaseModel, Field
import httpx
from tenacity import retry, stop_after_attempt, wait_exponential

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    NATIONAL_ID = "national_id"
    PASSPORT = "passport"
    DRIVERS_LICENSE = "drivers_license"
    VOTER_ID = "voter_id"
    UTILITY_BILL = "utility_bill"
    BANK_STATEMENT = "bank_statement"
    BUSINESS_REGISTRATION = "business_registration"
    TAX_CERTIFICATE = "tax_certificate"
    INVOICE = "invoice"
    RECEIPT = "receipt"
    CONTRACT = "contract"
    GENERIC = "generic"


class OCREngine(str, Enum):
    PADDLEOCR = "paddleocr"
    VLM = "vlm"
    DOCLING = "docling"
    TESSERACT = "tesseract"
    AUTO = "auto"


class ProcessingStatus(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"


@dataclass
class ServiceConfig:
    database_url: str = field(default_factory=lambda: os.getenv(
        "DATABASE_URL",
        "postgresql://postgres:postgres@localhost:5432/remittance"
    ))
    redis_url: str = field(default_factory=lambda: os.getenv("REDIS_URL", "redis://localhost:6379"))
    kafka_bootstrap_servers: str = field(default_factory=lambda: os.getenv("KAFKA_BOOTSTRAP_SERVERS", "localhost:9092"))
    lakehouse_url: str = field(default_factory=lambda: os.getenv("LAKEHOUSE_URL", "http://localhost:8181"))
    paddleocr_api_url: str = field(default_factory=lambda: os.getenv("PADDLEOCR_API_URL", "http://localhost:8024"))
    vlm_api_url: str = field(default_factory=lambda: os.getenv("VLM_API_URL", "http://localhost:8031"))
    vlm_api_key: str = field(default_factory=lambda: os.getenv("VLM_API_KEY", ""))
    docling_api_url: str = field(default_factory=lambda: os.getenv("DOCLING_API_URL", "http://localhost:8032"))
    document_storage_path: str = field(default_factory=lambda: os.getenv("DOCUMENT_STORAGE_PATH", "/tmp/ocr_documents"))
    max_file_size_mb: int = field(default_factory=lambda: int(os.getenv("MAX_FILE_SIZE_MB", "50")))


class DatabasePool:
    """Production-ready async database connection pool"""
    
    def __init__(self, database_url: str):
        self.database_url = database_url
        self._pool: Optional[asyncpg.Pool] = None
    
    async def initialize(self):
        if self._pool is None:
            self._pool = await asyncpg.create_pool(
                self.database_url,
                min_size=5,
                max_size=20,
                max_inactive_connection_lifetime=300,
                command_timeout=60
            )
            logger.info("Database pool initialized")
    
    async def close(self):
        if self._pool:
            await self._pool.close()
            self._pool = None
    
    @asynccontextmanager
    async def acquire(self) -> AsyncGenerator[asyncpg.Connection, None]:
        if self._pool is None:
            raise RuntimeError("Database pool not initialized")
        async with self._pool.acquire() as connection:
            yield connection
    
    @asynccontextmanager
    async def transaction(self) -> AsyncGenerator[asyncpg.Connection, None]:
        async with self.acquire() as connection:
            async with connection.transaction():
                yield connection


class RedisClient:
    """Production-ready Redis client"""
    
    def __init__(self, redis_url: str):
        self.redis_url = redis_url
        self._client: Optional[redis.Redis] = None
    
    async def initialize(self):
        if self._client is None:
            self._client = redis.from_url(
                self.redis_url,
                encoding="utf-8",
                decode_responses=True,
                max_connections=20
            )
            await self._client.ping()
            logger.info("Redis client initialized")
    
    async def close(self):
        if self._client:
            await self._client.close()
            self._client = None
    
    @property
    def client(self) -> redis.Redis:
        if self._client is None:
            raise RuntimeError("Redis client not initialized")
        return self._client


class KafkaProducer:
    """Kafka producer for event streaming"""
    
    def __init__(self, bootstrap_servers: str):
        self.bootstrap_servers = bootstrap_servers
        self._producer = None
    
    async def initialize(self):
        try:
            from aiokafka import AIOKafkaProducer
            self._producer = AIOKafkaProducer(
                bootstrap_servers=self.bootstrap_servers,
                value_serializer=lambda v: json.dumps(v).encode('utf-8'),
                key_serializer=lambda k: k.encode('utf-8') if k else None,
                acks='all'
            )
            await self._producer.start()
            logger.info("Kafka producer initialized")
        except ImportError:
            logger.warning("aiokafka not installed")
        except Exception as e:
            logger.warning(f"Kafka connection failed: {e}")
    
    async def close(self):
        if self._producer:
            await self._producer.stop()
            self._producer = None
    
    async def send_event(self, topic: str, key: str, value: Dict[str, Any]):
        if self._producer:
            try:
                await self._producer.send_and_wait(topic, value=value, key=key)
            except Exception as e:
                logger.error(f"Failed to send Kafka event: {e}")


class LakehouseClient:
    """Lakehouse client for analytics"""
    
    def __init__(self, url: str):
        self.url = url
        self._client: Optional[httpx.AsyncClient] = None
    
    async def initialize(self):
        self._client = httpx.AsyncClient(base_url=self.url, timeout=60.0)
        logger.info("Lakehouse client initialized")
    
    async def close(self):
        if self._client:
            await self._client.aclose()
            self._client = None
    
    async def write_event(self, table: str, data: Dict[str, Any]) -> bool:
        if not self._client:
            return True
        try:
            response = await self._client.post(f"/v1/tables/{table}/records", json=data)
            return response.status_code in (200, 201)
        except Exception as e:
            logger.error(f"Lakehouse write failed: {e}")
            return False


class PaddleOCRClient:
    """PaddleOCR client for fast text extraction with bounding boxes"""
    
    def __init__(self, api_url: str):
        self.api_url = api_url
        self._client: Optional[httpx.AsyncClient] = None
    
    async def initialize(self):
        self._client = httpx.AsyncClient(
            base_url=self.api_url,
            timeout=120.0
        )
        logger.info("PaddleOCR client initialized")
    
    async def close(self):
        if self._client:
            await self._client.aclose()
            self._client = None
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=1, max=10))
    async def extract_text(self, image_data: bytes, document_type: str) -> Dict[str, Any]:
        """Extract text from image using PaddleOCR"""
        if not self._client:
            return self._local_ocr_extraction(image_data, document_type)
        
        try:
            files = {"file": ("document.jpg", image_data)}
            data = {"document_type": document_type, "engines": "paddleocr"}
            response = await self._client.post("/ocr", files=files, data=data)
            
            if response.status_code == 200:
                result = response.json()
                text = result.get("text", "")
                confidence = result.get("confidence_score", 0.0)
                extracted_fields = self._extract_fields_from_text(text, document_type)
                return {
                    "engine": "paddleocr",
                    "text": text,
                    "confidence": confidence,
                    "document_type": document_type,
                    "extracted_fields": extracted_fields,
                    "tables": [],
                    "processed_at": datetime.utcnow().isoformat(),
                }
            
            logger.warning(f"PaddleOCR returned {response.status_code}, using local extraction")
            return self._local_ocr_extraction(image_data, document_type)
            
        except Exception as e:
            logger.warning(f"PaddleOCR failed: {e}, using local extraction")
            return self._local_ocr_extraction(image_data, document_type)
    
    def _local_ocr_extraction(self, image_data: bytes, document_type: str) -> Dict[str, Any]:
        """Local OCR extraction using Tesseract as fallback"""
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(io.BytesIO(image_data))
            
            text = pytesseract.image_to_string(image)
            
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            words = []
            confidences = []
            for i, word in enumerate(data['text']):
                if word.strip():
                    words.append({
                        "text": word,
                        "confidence": data['conf'][i] / 100.0 if data['conf'][i] > 0 else 0.5,
                        "bbox": {
                            "x": data['left'][i],
                            "y": data['top'][i],
                            "width": data['width'][i],
                            "height": data['height'][i]
                        }
                    })
                    if data['conf'][i] > 0:
                        confidences.append(data['conf'][i] / 100.0)
            
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.5
            
            extracted_fields = self._extract_fields_from_text(text, document_type)
            
            return {
                "engine": "tesseract",
                "text": text,
                "words": words,
                "confidence": avg_confidence,
                "document_type": document_type,
                "extracted_fields": extracted_fields,
                "tables": [],
                "processed_at": datetime.utcnow().isoformat()
            }
            
        except ImportError:
            logger.warning("pytesseract not installed, returning basic extraction")
            return {
                "engine": "basic",
                "text": "",
                "words": [],
                "confidence": 0.0,
                "document_type": document_type,
                "extracted_fields": {},
                "tables": [],
                "processed_at": datetime.utcnow().isoformat(),
                "error": "OCR libraries not available"
            }
        except Exception as e:
            logger.error(f"Local OCR extraction failed: {e}")
            return {
                "engine": "basic",
                "text": "",
                "words": [],
                "confidence": 0.0,
                "document_type": document_type,
                "extracted_fields": {},
                "tables": [],
                "processed_at": datetime.utcnow().isoformat(),
                "error": str(e)
            }
    
    def _extract_fields_from_text(self, text: str, document_type: str) -> Dict[str, Any]:
        """Extract structured fields from OCR text based on document type"""
        fields = {}
        
        nin_pattern = r'\b\d{11}\b'
        nin_matches = re.findall(nin_pattern, text)
        if nin_matches:
            fields["nin"] = nin_matches[0]
        
        date_pattern = r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})\b'
        date_matches = re.findall(date_pattern, text)
        if date_matches:
            fields["dates"] = date_matches
        
        name_pattern = r'(?:NAME|SURNAME|FIRST NAME|LAST NAME)[:\s]*([A-Z][a-zA-Z\s]+)'
        name_matches = re.findall(name_pattern, text, re.IGNORECASE)
        if name_matches:
            fields["names"] = [n.strip() for n in name_matches]
        
        if document_type == DocumentType.NATIONAL_ID.value:
            fields.update(self._extract_national_id_fields(text))
        elif document_type == DocumentType.PASSPORT.value:
            fields.update(self._extract_passport_fields(text))
        elif document_type == DocumentType.BUSINESS_REGISTRATION.value:
            fields.update(self._extract_business_reg_fields(text))
        elif document_type == DocumentType.UTILITY_BILL.value:
            fields.update(self._extract_utility_bill_fields(text))
        elif document_type == DocumentType.BANK_STATEMENT.value:
            fields.update(self._extract_bank_statement_fields(text))
        
        return fields
    
    def _extract_national_id_fields(self, text: str) -> Dict[str, Any]:
        """Extract fields specific to Nigerian National ID"""
        fields = {}
        
        nin_pattern = r'(?:NIN|NATIONAL IDENTIFICATION NUMBER)[:\s]*(\d{11})'
        nin_match = re.search(nin_pattern, text, re.IGNORECASE)
        if nin_match:
            fields["nin"] = nin_match.group(1)
        
        dob_pattern = r'(?:DATE OF BIRTH|DOB|BIRTH DATE)[:\s]*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'
        dob_match = re.search(dob_pattern, text, re.IGNORECASE)
        if dob_match:
            fields["date_of_birth"] = dob_match.group(1)
        
        gender_pattern = r'(?:SEX|GENDER)[:\s]*(MALE|FEMALE|M|F)'
        gender_match = re.search(gender_pattern, text, re.IGNORECASE)
        if gender_match:
            fields["gender"] = gender_match.group(1).upper()
        
        return fields
    
    def _extract_passport_fields(self, text: str) -> Dict[str, Any]:
        """Extract fields specific to passport"""
        fields = {}
        
        passport_pattern = r'(?:PASSPORT NO|PASSPORT NUMBER)[:\s]*([A-Z]\d{8})'
        passport_match = re.search(passport_pattern, text, re.IGNORECASE)
        if passport_match:
            fields["passport_number"] = passport_match.group(1)
        
        mrz_pattern = r'P<[A-Z]{3}[A-Z<]+<<[A-Z<]+'
        mrz_match = re.search(mrz_pattern, text)
        if mrz_match:
            fields["mrz_line1"] = mrz_match.group(0)
        
        return fields
    
    def _extract_business_reg_fields(self, text: str) -> Dict[str, Any]:
        """Extract fields specific to business registration"""
        fields = {}
        
        cac_pattern = r'(?:RC|BN|IT)\s*(\d{6,8})'
        cac_match = re.search(cac_pattern, text, re.IGNORECASE)
        if cac_match:
            fields["registration_number"] = cac_match.group(0).replace(" ", "")
        
        tin_pattern = r'(\d{8}-\d{4})'
        tin_match = re.search(tin_pattern, text)
        if tin_match:
            fields["tax_id"] = tin_match.group(1)
        
        return fields
    
    def _extract_utility_bill_fields(self, text: str) -> Dict[str, Any]:
        """Extract fields specific to utility bills"""
        fields = {}
        
        amount_pattern = r'(?:AMOUNT|TOTAL|DUE)[:\s]*(?:NGN|₦)?\s*([\d,]+\.?\d*)'
        amount_match = re.search(amount_pattern, text, re.IGNORECASE)
        if amount_match:
            fields["amount"] = amount_match.group(1).replace(",", "")
        
        account_pattern = r'(?:ACCOUNT|METER)[:\s]*(\d{10,15})'
        account_match = re.search(account_pattern, text, re.IGNORECASE)
        if account_match:
            fields["account_number"] = account_match.group(1)
        
        return fields
    
    def _extract_bank_statement_fields(self, text: str) -> Dict[str, Any]:
        """Extract fields specific to bank statements"""
        fields = {}
        
        account_pattern = r'(?:ACCOUNT NUMBER|ACCT NO)[:\s]*(\d{10})'
        account_match = re.search(account_pattern, text, re.IGNORECASE)
        if account_match:
            fields["account_number"] = account_match.group(1)
        
        balance_pattern = r'(?:CLOSING BALANCE|BALANCE)[:\s]*(?:NGN|₦)?\s*([\d,]+\.?\d*)'
        balance_match = re.search(balance_pattern, text, re.IGNORECASE)
        if balance_match:
            fields["closing_balance"] = balance_match.group(1).replace(",", "")
        
        return fields


class VLMClient:
    """Vision Language Model client for semantic document understanding"""
    
    def __init__(self, api_url: str, api_key: str):
        self.api_url = api_url
        self.api_key = api_key
        self._client: Optional[httpx.AsyncClient] = None
    
    async def initialize(self):
        headers = {}
        if self.api_key:
            headers["Authorization"] = f"Bearer {self.api_key}"
        self._client = httpx.AsyncClient(
            base_url=self.api_url,
            headers=headers,
            timeout=120.0
        )
        logger.info("VLM client initialized")
    
    async def close(self):
        if self._client:
            await self._client.aclose()
            self._client = None
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=1, max=10))
    async def extract_text(self, image_data: bytes, document_type: str) -> Dict[str, Any]:
        """Analyse document image using a Vision Language Model."""
        if not self._client:
            return {
                "engine": "vlm",
                "text": "",
                "confidence": 0.0,
                "document_type": document_type,
                "extracted_fields": {},
                "semantic_labels": {},
                "tables": [],
                "processed_at": datetime.utcnow().isoformat(),
                "error": "VLM client not available",
            }
        try:
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            response = await self._client.post(
                "/v1/ocr/extract",
                json={
                    "image": image_base64,
                    "document_type": document_type,
                    "language": "en",
                    "extract_tables": True,
                    "extract_fields": True,
                },
            )
            if response.status_code == 200:
                result = response.json()
                return {
                    "engine": "vlm",
                    "text": result.get("text", ""),
                    "confidence": result.get("confidence", 0.0),
                    "document_type": document_type,
                    "extracted_fields": result.get("extracted_fields", {}),
                    "semantic_labels": result.get("semantic_labels", {}),
                    "tables": result.get("tables", []),
                    "processed_at": datetime.utcnow().isoformat(),
                }
            logger.warning(f"VLM returned {response.status_code}")
        except Exception as e:
            logger.warning(f"VLM failed: {e}")
        return {
            "engine": "vlm",
            "text": "",
            "confidence": 0.0,
            "document_type": document_type,
            "extracted_fields": {},
            "semantic_labels": {},
            "tables": [],
            "processed_at": datetime.utcnow().isoformat(),
        }


class DoclingClient:
    """Docling client for document understanding and structured extraction"""
    
    def __init__(self, api_url: str):
        self.api_url = api_url
        self._client: Optional[httpx.AsyncClient] = None
    
    async def initialize(self):
        self._client = httpx.AsyncClient(
            base_url=self.api_url,
            timeout=180.0
        )
        logger.info("Docling client initialized")
    
    async def close(self):
        if self._client:
            await self._client.aclose()
            self._client = None
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=1, max=10))
    async def process_document(self, file_data: bytes, file_name: str, document_type: str) -> Dict[str, Any]:
        """Process document using Docling for structured extraction"""
        if not self._client:
            return self._local_document_processing(file_data, file_name, document_type)
        
        try:
            files = {"file": (file_name, file_data)}
            data = {"document_type": document_type}
            
            response = await self._client.post(
                "/v1/documents/process",
                files=files,
                data=data
            )
            
            if response.status_code == 200:
                return response.json()
            
            logger.warning(f"Docling returned {response.status_code}, using local processing")
            return self._local_document_processing(file_data, file_name, document_type)
            
        except Exception as e:
            logger.warning(f"Docling processing failed: {e}, using local processing")
            return self._local_document_processing(file_data, file_name, document_type)
    
    def _local_document_processing(self, file_data: bytes, file_name: str, document_type: str) -> Dict[str, Any]:
        """Local document processing fallback"""
        try:
            file_ext = Path(file_name).suffix.lower()
            
            if file_ext == '.pdf':
                return self._process_pdf(file_data, document_type)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                return self._process_image(file_data, document_type)
            elif file_ext in ['.doc', '.docx']:
                return self._process_word(file_data, document_type)
            else:
                return {
                    "engine": "local",
                    "file_name": file_name,
                    "document_type": document_type,
                    "content": "",
                    "structure": {},
                    "tables": [],
                    "metadata": {},
                    "processed_at": datetime.utcnow().isoformat(),
                    "error": f"Unsupported file type: {file_ext}"
                }
                
        except Exception as e:
            logger.error(f"Local document processing failed: {e}")
            return {
                "engine": "local",
                "file_name": file_name,
                "document_type": document_type,
                "content": "",
                "structure": {},
                "tables": [],
                "metadata": {},
                "processed_at": datetime.utcnow().isoformat(),
                "error": str(e)
            }
    
    def _process_pdf(self, file_data: bytes, document_type: str) -> Dict[str, Any]:
        """Process PDF document"""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(io.BytesIO(file_data))
            
            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text() or "")
            
            full_text = "\n\n".join(text_content)
            
            metadata = {}
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                    "creation_date": str(reader.metadata.get("/CreationDate", "")),
                }
            
            return {
                "engine": "pypdf",
                "document_type": document_type,
                "content": full_text,
                "page_count": len(reader.pages),
                "structure": {
                    "pages": [{"page_number": i+1, "text": t} for i, t in enumerate(text_content)]
                },
                "tables": [],
                "metadata": metadata,
                "processed_at": datetime.utcnow().isoformat()
            }
            
        except ImportError:
            logger.warning("pypdf not installed")
            return {
                "engine": "local",
                "document_type": document_type,
                "content": "",
                "structure": {},
                "tables": [],
                "metadata": {},
                "processed_at": datetime.utcnow().isoformat(),
                "error": "PDF processing library not available"
            }
    
    def _process_image(self, file_data: bytes, document_type: str) -> Dict[str, Any]:
        """Process image document"""
        try:
            from PIL import Image
            
            image = Image.open(io.BytesIO(file_data))
            
            return {
                "engine": "pillow",
                "document_type": document_type,
                "content": "",
                "structure": {},
                "tables": [],
                "metadata": {
                    "format": image.format,
                    "mode": image.mode,
                    "size": {"width": image.width, "height": image.height}
                },
                "processed_at": datetime.utcnow().isoformat(),
                "note": "Image loaded successfully, use OCR endpoint for text extraction"
            }
            
        except ImportError:
            return {
                "engine": "local",
                "document_type": document_type,
                "content": "",
                "structure": {},
                "tables": [],
                "metadata": {},
                "processed_at": datetime.utcnow().isoformat(),
                "error": "Image processing library not available"
            }
    
    def _process_word(self, file_data: bytes, document_type: str) -> Dict[str, Any]:
        """Process Word document"""
        try:
            import docx
            
            doc = docx.Document(io.BytesIO(file_data))
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "engine": "python-docx",
                "document_type": document_type,
                "content": full_text,
                "structure": {
                    "paragraphs": paragraphs
                },
                "tables": tables,
                "metadata": {
                    "paragraph_count": len(paragraphs),
                    "table_count": len(tables)
                },
                "processed_at": datetime.utcnow().isoformat()
            }
            
        except ImportError:
            return {
                "engine": "local",
                "document_type": document_type,
                "content": "",
                "structure": {},
                "tables": [],
                "metadata": {},
                "processed_at": datetime.utcnow().isoformat(),
                "error": "Word processing library not available"
            }


class DocumentValidator:
    """Document validation and quality assessment"""
    
    @classmethod
    def validate_document(cls, ocr_result: Dict[str, Any], document_type: str) -> Dict[str, Any]:
        """Validate OCR result and assess document quality"""
        
        validation_result = {
            "is_valid": True,
            "confidence_score": ocr_result.get("confidence", 0.0),
            "quality_score": 0.0,
            "issues": [],
            "warnings": [],
            "extracted_fields_valid": True
        }
        
        confidence = ocr_result.get("confidence", 0.0)
        if confidence < 0.3:
            validation_result["is_valid"] = False
            validation_result["issues"].append("Very low OCR confidence - document may be unreadable")
        elif confidence < 0.6:
            validation_result["warnings"].append("Low OCR confidence - some text may be incorrect")
        
        text = ocr_result.get("text", "")
        if len(text) < 50:
            validation_result["warnings"].append("Very little text extracted - document may be mostly images")
        
        quality_score = cls._calculate_quality_score(ocr_result)
        validation_result["quality_score"] = quality_score
        
        if quality_score < 0.3:
            validation_result["is_valid"] = False
            validation_result["issues"].append("Document quality too low for reliable extraction")
        
        fields_validation = cls._validate_extracted_fields(
            ocr_result.get("extracted_fields", {}),
            document_type
        )
        validation_result["fields_validation"] = fields_validation
        if not fields_validation.get("all_required_present", True):
            validation_result["warnings"].append("Some required fields could not be extracted")
        
        return validation_result
    
    @classmethod
    def _calculate_quality_score(cls, ocr_result: Dict[str, Any]) -> float:
        """Calculate overall document quality score"""
        
        scores = []
        
        confidence = ocr_result.get("confidence", 0.0)
        scores.append(confidence)
        
        words = ocr_result.get("words", [])
        if words:
            word_confidences = [w.get("confidence", 0.5) for w in words]
            avg_word_confidence = sum(word_confidences) / len(word_confidences)
            scores.append(avg_word_confidence)
        
        text = ocr_result.get("text", "")
        if text:
            readable_chars = sum(1 for c in text if c.isalnum() or c.isspace())
            total_chars = len(text)
            readability = readable_chars / total_chars if total_chars > 0 else 0
            scores.append(readability)
        
        return sum(scores) / len(scores) if scores else 0.0
    
    @classmethod
    def _validate_extracted_fields(cls, fields: Dict[str, Any], document_type: str) -> Dict[str, Any]:
        """Validate extracted fields based on document type"""
        
        required_fields = {
            DocumentType.NATIONAL_ID.value: ["nin"],
            DocumentType.PASSPORT.value: ["passport_number"],
            DocumentType.BUSINESS_REGISTRATION.value: ["registration_number"],
            DocumentType.UTILITY_BILL.value: ["account_number", "amount"],
            DocumentType.BANK_STATEMENT.value: ["account_number"]
        }
        
        required = required_fields.get(document_type, [])
        present = [f for f in required if f in fields and fields[f]]
        missing = [f for f in required if f not in fields or not fields[f]]
        
        return {
            "required_fields": required,
            "present_fields": present,
            "missing_fields": missing,
            "all_required_present": len(missing) == 0
        }


class OCRRequest(BaseModel):
    document_type: DocumentType = DocumentType.GENERIC
    engine: OCREngine = OCREngine.AUTO
    extract_tables: bool = True
    extract_fields: bool = True
    language: str = "en"


class OCRResponse(BaseModel):
    job_id: str
    status: str
    document_type: str
    engine_used: str
    text: Optional[str] = None
    confidence: Optional[float] = None
    extracted_fields: Dict[str, Any] = {}
    tables: List[Any] = []
    validation: Optional[Dict[str, Any]] = None
    processing_time_ms: Optional[int] = None
    created_at: datetime


class ServiceContainer:
    """Container for all service dependencies"""
    
    def __init__(self, config: ServiceConfig):
        self.config = config
        self.db = DatabasePool(config.database_url)
        self.redis = RedisClient(config.redis_url)
        self.kafka = KafkaProducer(config.kafka_bootstrap_servers)
        self.lakehouse = LakehouseClient(config.lakehouse_url)
        self.paddleocr = PaddleOCRClient(config.paddleocr_api_url)
        self.vlm = VLMClient(config.vlm_api_url, config.vlm_api_key)
        self.docling = DoclingClient(config.docling_api_url)
    
    async def initialize(self):
        await self.db.initialize()
        await self.redis.initialize()
        await self.kafka.initialize()
        await self.lakehouse.initialize()
        await self.paddleocr.initialize()
        await self.vlm.initialize()
        await self.docling.initialize()
        await self._ensure_tables()
        logger.info("All services initialized")
    
    async def close(self):
        await self.docling.close()
        await self.vlm.close()
        await self.paddleocr.close()
        await self.lakehouse.close()
        await self.kafka.close()
        await self.redis.close()
        await self.db.close()
        logger.info("All services closed")
    
    async def _ensure_tables(self):
        """Ensure all required tables exist"""
        async with self.db.acquire() as conn:
            await conn.execute("""
                CREATE TABLE IF NOT EXISTS ocr_jobs (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    file_name VARCHAR(255),
                    file_hash VARCHAR(64),
                    file_size INTEGER,
                    mime_type VARCHAR(100),
                    document_type VARCHAR(50),
                    engine_used VARCHAR(50),
                    status VARCHAR(50) DEFAULT 'pending',
                    ocr_result JSONB,
                    validation_result JSONB,
                    confidence DECIMAL(5,4),
                    processing_time_ms INTEGER,
                    error_message TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    completed_at TIMESTAMP
                )
            """)
            
            await conn.execute("""
                CREATE INDEX IF NOT EXISTS idx_ocr_jobs_file_hash ON ocr_jobs(file_hash)
            """)
            
            await conn.execute("""
                CREATE INDEX IF NOT EXISTS idx_ocr_jobs_status ON ocr_jobs(status)
            """)
            
            logger.info("Database tables ensured")


services: Optional[ServiceContainer] = None


def get_services() -> ServiceContainer:
    if services is None:
        raise RuntimeError("Services not initialized")
    return services


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan manager"""
    global services
    
    try:
        config = ServiceConfig()
        services = ServiceContainer(config)
        await services.initialize()
        yield
    finally:
        if services:
            await services.close()


app = FastAPI(
    title="OCR Service (Production)",
    description="Production-ready OCR service with PaddleOCR, VLM, and Docling integration",
    version="2.0.0",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("ALLOWED_ORIGINS","http://localhost:5173,http://localhost:5174,http://localhost:3000").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.post("/ocr/extract", response_model=OCRResponse)
async def extract_text_from_document(
    file: UploadFile = File(...),
    document_type: DocumentType = Form(DocumentType.GENERIC),
    engine: OCREngine = Form(OCREngine.AUTO),
    extract_tables: bool = Form(True),
    extract_fields: bool = Form(True),
    svc: ServiceContainer = Depends(get_services)
):
    """Extract text from uploaded document using OCR"""
    
    start_time = datetime.utcnow()
    
    file_content = await file.read()
    file_hash = hashlib.sha256(file_content).hexdigest()
    file_size = len(file_content)
    
    max_size = svc.config.max_file_size_mb * 1024 * 1024
    if file_size > max_size:
        raise HTTPException(
            status_code=400,
            detail=f"File size exceeds maximum allowed ({svc.config.max_file_size_mb}MB)"
        )
    
    cached_result = await svc.redis.client.get(f"ocr:{file_hash}")
    if cached_result:
        cached_data = json.loads(cached_result)
        return OCRResponse(**cached_data)
    
    async with svc.db.transaction() as conn:
        job = await conn.fetchrow(
            """
            INSERT INTO ocr_jobs (file_name, file_hash, file_size, mime_type, document_type, status)
            VALUES ($1, $2, $3, $4, $5, $6)
            RETURNING *
            """,
            file.filename, file_hash, file_size, file.content_type,
            document_type.value, ProcessingStatus.PROCESSING.value
        )
    
    try:
        if engine == OCREngine.AUTO:
            file_ext = Path(file.filename).suffix.lower()
            if file_ext in ['.pdf', '.doc', '.docx']:
                engine = OCREngine.DOCLING
            else:
                engine = OCREngine.PADDLEOCR
        
        if engine == OCREngine.PADDLEOCR:
            ocr_result = await svc.paddleocr.extract_text(file_content, document_type.value)
            engine_used = ocr_result.get("engine", "paddleocr")
        elif engine == OCREngine.VLM:
            ocr_result = await svc.vlm.extract_text(file_content, document_type.value)
            engine_used = ocr_result.get("engine", "vlm")
        elif engine == OCREngine.DOCLING:
            ocr_result = await svc.docling.process_document(
                file_content, file.filename, document_type.value
            )
            engine_used = ocr_result.get("engine", "docling")
        else:
            ocr_result = await svc.paddleocr.extract_text(file_content, document_type.value)
            engine_used = ocr_result.get("engine", "paddleocr")
        
        validation_result = DocumentValidator.validate_document(ocr_result, document_type.value)
        
        end_time = datetime.utcnow()
        processing_time_ms = int((end_time - start_time).total_seconds() * 1000)
        
        async with svc.db.transaction() as conn:
            await conn.execute(
                """
                UPDATE ocr_jobs
                SET status = $1, engine_used = $2, ocr_result = $3, validation_result = $4,
                    confidence = $5, processing_time_ms = $6, completed_at = $7
                WHERE id = $8
                """,
                ProcessingStatus.COMPLETED.value, engine_used,
                json.dumps(ocr_result), json.dumps(validation_result),
                ocr_result.get("confidence", 0.0), processing_time_ms,
                datetime.utcnow(), job['id']
            )
        
        response = OCRResponse(
            job_id=str(job['id']),
            status=ProcessingStatus.COMPLETED.value,
            document_type=document_type.value,
            engine_used=engine_used,
            text=ocr_result.get("text", ""),
            confidence=ocr_result.get("confidence"),
            extracted_fields=ocr_result.get("extracted_fields", {}),
            tables=ocr_result.get("tables", []),
            validation=validation_result,
            processing_time_ms=processing_time_ms,
            created_at=job['created_at']
        )
        
        await svc.redis.client.setex(
            f"ocr:{file_hash}",
            3600,
            json.dumps(response.dict(), default=str)
        )
        
        event_data = {
            "event_type": "ocr.document_processed",
            "job_id": str(job['id']),
            "document_type": document_type.value,
            "engine_used": engine_used,
            "confidence": ocr_result.get("confidence"),
            "processing_time_ms": processing_time_ms,
            "timestamp": datetime.utcnow().isoformat()
        }
        await svc.kafka.send_event("ocr-events", str(job['id']), event_data)
        await svc.lakehouse.write_event("ocr_events", event_data)
        
        return response
        
    except Exception as e:
        logger.error(f"OCR processing failed: {e}")
        
        async with svc.db.transaction() as conn:
            await conn.execute(
                """
                UPDATE ocr_jobs
                SET status = $1, error_message = $2, completed_at = $3
                WHERE id = $4
                """,
                ProcessingStatus.FAILED.value, str(e),
                datetime.utcnow(), job['id']
            )
        
        raise HTTPException(status_code=500, detail=f"OCR processing failed: {str(e)}")


@app.post("/documents/process")
async def process_document(
    file: UploadFile = File(...),
    document_type: DocumentType = Form(DocumentType.GENERIC),
    svc: ServiceContainer = Depends(get_services)
):
    """Process document using Docling for structured extraction"""
    
    file_content = await file.read()
    
    result = await svc.docling.process_document(
        file_content, file.filename, document_type.value
    )
    
    return {
        "file_name": file.filename,
        "document_type": document_type.value,
        "result": result
    }


@app.get("/ocr/jobs/{job_id}")
async def get_ocr_job(
    job_id: str,
    svc: ServiceContainer = Depends(get_services)
):
    """Get OCR job status and result"""
    
    async with svc.db.acquire() as conn:
        result = await conn.fetchrow(
            "SELECT * FROM ocr_jobs WHERE id = $1",
            uuid.UUID(job_id)
        )
        if not result:
            raise HTTPException(status_code=404, detail="Job not found")
        
        return {
            "job_id": str(result['id']),
            "file_name": result['file_name'],
            "document_type": result['document_type'],
            "engine_used": result['engine_used'],
            "status": result['status'],
            "confidence": float(result['confidence']) if result['confidence'] else None,
            "processing_time_ms": result['processing_time_ms'],
            "ocr_result": json.loads(result['ocr_result']) if result['ocr_result'] else None,
            "validation_result": json.loads(result['validation_result']) if result['validation_result'] else None,
            "error_message": result['error_message'],
            "created_at": result['created_at'].isoformat(),
            "completed_at": result['completed_at'].isoformat() if result['completed_at'] else None
        }


@app.get("/ocr/supported-types")
async def get_supported_document_types():
    """Get list of supported document types"""
    
    return {
        "document_types": [
            {"value": dt.value, "name": dt.name.replace("_", " ").title()}
            for dt in DocumentType
        ],
        "engines": [
            {"value": e.value, "name": e.name.title(), "description": _get_engine_description(e)}
            for e in OCREngine
        ],
        "supported_formats": [
            ".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".gif",
            ".pdf", ".doc", ".docx"
        ]
    }


def _get_engine_description(engine: OCREngine) -> str:
    """Get description for OCR engine"""
    descriptions = {
        OCREngine.PADDLEOCR: "PaddleOCR - Fast text extraction with bounding boxes",
        OCREngine.VLM: "Vision Language Model - Semantic document understanding",
        OCREngine.DOCLING: "Docling - Structured document parsing and layout analysis",
        OCREngine.TESSERACT: "Tesseract - Open source OCR fallback",
        OCREngine.AUTO: "Automatic - Selects best engine based on file type",
    }
    return descriptions.get(engine, "")


@app.get("/health")
async def health_check(svc: ServiceContainer = Depends(get_services)):
    """Health check endpoint"""
    
    health_status = {
        "status": "healthy",
        "service": "OCR Service (Production)",
        "version": "2.0.0",
        "timestamp": datetime.utcnow().isoformat(),
        "components": {}
    }
    
    try:
        async with svc.db.acquire() as conn:
            await conn.fetchval("SELECT 1")
        health_status["components"]["database"] = "healthy"
    except Exception as e:
        health_status["components"]["database"] = f"unhealthy: {str(e)}"
        health_status["status"] = "degraded"
    
    try:
        await svc.redis.client.ping()
        health_status["components"]["redis"] = "healthy"
    except Exception as e:
        health_status["components"]["redis"] = f"unhealthy: {str(e)}"
        health_status["status"] = "degraded"
    
    return health_status


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8030)
